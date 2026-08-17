"""Generate the Graxella Project Charter (Word document).

Run:  C:\\Python313\\python.exe build_project_charter.py
Output: Graxella_Project_Charter.docx (in this folder).

The charter is the durable reference for the program: problem, solution,
differentiators, competitive relevance, target audience, metrics, roadmap,
scope, principles, components, risks, glossary.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

GREEN = RGBColor(0x14, 0x80, 0x5E)
INK = RGBColor(0x1C, 0x25, 0x21)
MUTED = RGBColor(0x5C, 0x6A, 0x63)
AMBER = RGBColor(0xA8, 0x69, 0x1C)

doc = Document()

# ---- base styles ----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for level, size in ((1, 17), (2, 13.5), (3, 11.5)):
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Segoe UI"
    st.font.bold = True
    st.font.size = Pt(size)
    st.font.color.rgb = GREEN if level > 1 else INK
    st.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    st.paragraph_format.space_after = Pt(6)
    # keep Word from theming the color back
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")


def p(text="", *, bold=False, italic=False, color=None, size=None, style=None):
    para = doc.add_paragraph(style=style)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    return para


def lead(head, body):
    para = doc.add_paragraph()
    r1 = para.add_run(head + " — ")
    r1.bold = True
    r1.font.color.rgb = GREEN
    para.add_run(body)
    return para


def bullets(items, *, style="List Bullet"):
    for item in items:
        if isinstance(item, tuple):
            head, body = item
            para = doc.add_paragraph(style=style)
            r1 = para.add_run(head + " — ")
            r1.bold = True
            r1.font.color.rgb = GREEN
            para.add_run(body)
        else:
            doc.add_paragraph(item, style=style)


def table(rows, *, widths=None, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = t.cell(ri, ci)
            if widths:
                cell.width = Inches(widths[ci])
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            if header and ri == 0:
                run.bold = True
                run.font.color.rgb = GREEN
                run.font.name = "Segoe UI"
                run.font.size = Pt(9)
            para.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()
    return t


# ===========================================================================
# COVER
# ===========================================================================
title = p("GRAXELLA", bold=True, size=34)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = p("Project Charter — The Evidence Loop", size=16, color=GREEN, bold=True)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag = p("Reliable AI agents at enterprise scale: a runtime where reliability compounds "
        "with volume, and every behavior change ships with citations.",
        italic=True, color=MUTED, size=11)
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
table([
    ["Document", "Graxella Project Charter"],
    ["Version / Status", "v1.0 (Draft) — living reference; re-issue on phase completion"],
    ["Date", "17 August 2026"],
    ["Program scope", "graxella (umbrella) · mnema · agent2society · axon (B8/B10 tool fabric)"],
    ["Companion documents", "The Evidence Loop plan · Research Radar · MAST Coverage Scorecard · Flagship deck (PPTX)"],
], widths=[1.8, 4.6], header=False)

doc.add_page_break()

# ===========================================================================
# 1. EXECUTIVE SUMMARY
# ===========================================================================
doc.add_heading("1. Executive Summary", level=1)
p("Enterprises can build AI agents; they cannot yet trust them. Multi-step agent "
  "workflows fail in production because per-step errors compound, fixes live in "
  "unversioned prompt edits, audits cannot explain decisions, and — most "
  "fundamentally — nothing in today's stacks converts operational volume into "
  "reliability. Month six of an agent deployment is exactly as fragile as week one.")
p("Graxella is an evidence-based reliability and governance layer — a control "
  "plane — that runs underneath the orchestration frameworks enterprises already "
  "use (LangGraph, LangChain, CrewAI, AutoGen, MCP, A2A). It wraps existing agents "
  "in one line of code, records every decision and outcome as typed, immutable, "
  "cited evidence, mines that evidence offline for improvements, and promotes "
  "fixes through a Bayesian gate that requires proof before any behavior change "
  "ships. The runtime hot path stays deterministic and LLM-free; all learning "
  "happens at compile time.")
p("The one-sentence promise: agents that get more reliable the more you run "
  "them — with a citation for every change and a rollback for every mistake.")

# ===========================================================================
# 2. VISION AND MISSION
# ===========================================================================
doc.add_heading("2. Vision and Mission", level=1)
lead("Vision", "Agentic AI becomes dependable infrastructure: enterprises operate "
     "fleets of agents with the same confidence, audit discipline, and compounding "
     "improvement they expect from mature software systems.")
lead("Mission", "Build the substrate that makes any agent stack trustworthy — an "
     "'optimistic orchestration' layer where the flow the developer wrote always "
     "runs exactly as written, while underneath, the system compresses what goes "
     "in, heals what breaks, routes without deliberation, remembers everything, "
     "and learns from every failure so nothing fails the same way twice.")
lead("The four-pillar bar", "Every design decision must clear complexity (scales "
     "to thousands of agents and domains without new rules), clarity (every number "
     "traceable to named evidence), optimization (self-calibrating from outcomes), "
     "and reliability (deterministic, O(1) governance, no LLM in any decision loop).")

# ===========================================================================
# 3. THE PROBLEM
# ===========================================================================
doc.add_heading("3. The Problem We Are Solving", level=1)

doc.add_heading("3.1 The reliability wall", level=2)
p("Agent chains multiply per-step error: a 95%-reliable step compounds to roughly "
  "60% success over a ten-step workflow; 90% per-step collapses to roughly 35%. "
  "Demos are one friendly step; production business processes are ten steps across "
  "systems that drift weekly. This is why the industry's agent programs stall in "
  "'pilot purgatory'.")
p("This is now peer-reviewed fact, not anecdote. MAST (UC Berkeley, arXiv "
  "2503.13657) analyzed 1,600+ annotated failure traces across seven frameworks "
  "and identified 14 failure modes in three clusters — system design issues "
  "(44.2% of failures), inter-agent misalignment (32.3%), and task verification "
  "(23.5%). Its headline conclusion is graxella's thesis: most failures stem from "
  "orchestration design, not model weakness. Bigger models do not fix this; "
  "better substrates do.")

doc.add_heading("3.2 The five value leaks", level=2)
p("Enterprises leak value at five specific points in every agent deployment. "
  "These five leaks define graxella's five capability stages (Section 4.3).")
bullets([
    ("Leak 1 · Input tokens", "Prompts and context are assembled wastefully; "
     "retries burn tokens silently; cache-hostile prompt layouts forfeit up to "
     "90% provider discounts. Cost per completed task is opaque to everyone, "
     "including the CFO who pays for it."),
    ("Leak 2 · The tool boundary", "Upstream APIs drift — fields rename, schemas "
     "version, tools deprecate. Agents keep calling the old shape, burn retries, "
     "and every incident ends with a human editing a prompt. Most production "
     "drift originates here."),
    ("Leak 3 · Agent coordination", "Supervisor LLMs re-read every agent's full "
     "description on every hop ('waiting tokens'), then guess. Agents do not know "
     "each other's limitations; handoffs are informal; coordination failures are "
     "32.3% of observed failure mass (MAST)."),
    ("Leak 4 · Memory", "State lives in context windows and chat scrollback. "
     "History is lost, contradictions accumulate invisibly, and 'what did we "
     "believe when we made that decision?' is unanswerable."),
    ("Leak 5 · Learning", "The most expensive leak: every run produces evidence "
     "about what works, and today's stacks throw it away as logs. Operations do "
     "not compound; ops headcount scales linearly with agent count."),
])

doc.add_heading("3.3 Why current stacks cannot fix it", level=2)
bullets([
    ("Runtime-LLM answers get worse with scale", "LLM judges, critics, and "
     "retry-with-a-smarter-prompt add cost per call, forever, and are "
     "non-deterministic and un-auditable. Per-decision LLM governance is exactly "
     "wrong when volume is the goal."),
    ("Fixes are prose", "Prompt edits are unversioned, unreviewed, unmeasured, "
     "and unrollbackable. Agent behavior has no change management."),
    ("Audit is forensics", "Trace viewers show what happened; they cannot say "
     "why, because the 'why' was never recorded. 'The LLM decided' fails EU AI "
     "Act, model-risk-management, and SOX conversations."),
    ("Protocols move messages, not trust", "MCP, A2A, and ACP standardize "
     "transport and discovery, but — as the 2026 'Governance Gaps' analysis "
     "(arXiv 2606.31498) shows — they cannot express governance, accountability, "
     "or evidence of capability. The layer above them is empty."),
])

# ===========================================================================
# 4. HOW WE ARE SOLVING IT
# ===========================================================================
doc.add_heading("4. How We Are Solving It", level=1)

doc.add_heading("4.1 The thesis: behavior as an evidence-promoted artifact", level=2)
p("Graxella inverts the industry default. Instead of adding runtime LLM judgment, "
  "it moves ALL learning to compile time and keeps the hot path deterministic. "
  "The unit of improvement is not a prompt edit but a promotion: a versioned, "
  "rollbackable behavior change that carries its own citations into an immutable "
  "ledger. The analogy is GitOps — with the pull-request reviewer replaced by a "
  "Bayesian gate that reads the organization's own operational history.")
p("We call the operating doctrine 'optimistic orchestration': the flow the "
  "developer wrote always runs, exactly as written (protected by hard "
  "constitution invariants that no evidence can override). Learning never blocks "
  "and never mutates the flow topology; it only adjusts within safe envelopes — "
  "prompts, transforms, tool bindings, routing weights — each change gated, "
  "cited, and reversible.")

doc.add_heading("4.2 The Evidence Loop", level=2)
p("Six stages, one substrate:")
bullets([
    ("1 · Dispatch", "Deterministic, O(1), zero runtime LLM. Tasks route to "
     "agents via auditable scoring (matched tokens and tags recorded per "
     "decision); tool calls apply promoted healing transforms."),
    ("2 · Outcome", "Every dispatch auto-records a typed outcome (success, "
     "failure, latency, cost, violations) linked to its decision — per hop and "
     "per trajectory. Verification is structural, not opt-in."),
    ("3 · Ledger", "Mnema stores decisions and outcomes as immutable, "
     "bi-temporal, provenance-mandatory assertions. History cannot be rewritten "
     "by construction."),
    ("4 · Proposal", "Offline miners (drift patterns, docs, DSPy/GEPA prompt "
     "compilation, paired-trajectory audits) read the ledger and propose concrete "
     "changes. LLMs are permitted here — it is compile time."),
    ("5 · Gate", "A proposal ships only when the Bayesian gate finds sufficient "
     "cited history for that exact (domain, kind, model) tuple — or a human signs "
     "off. Cold domains default to human review; thresholds self-calibrate per "
     "tuple; provenance diversity (≥K independent sessions) defends against "
     "evidence poisoning. Zero LLM in the decision."),
    ("6 · Promotion", "Approved changes ship as versioned artifacts with rollback "
     "points. The next dispatch reads promoted behavior — reliability compounds "
     "with volume."),
])

doc.add_heading("4.3 The five capability stages (founding vision)", level=2)
table([
    ["Stage", "Capability", "Key mechanisms"],
    ["1 · Token economy", "Reduce input cost without changing intent",
     "Cache-aware prompt layout; token accounting in the value ledger; "
     "DSPy/GEPA-compiled prompts as gated promotions; extractive compression "
     "(LLMLingua-2 class) only for long novel content, evidence-gated per domain"],
    ["2 · Tool boundary", "Tool failures and drift healed; correct usage learned",
     "Knowledge graph of tools with live cited trust scores; drift interception "
     "(B10 middleware); healer→compiler→verifier producing deterministic "
     "transforms; sentinel canaries/shadow calls for fail-zero predictive drift "
     "detection; 'fail once, learn forever'"],
    ["3 · Agent-to-agent", "Deterministic coordination; no waiting tokens",
     "agent2society TF-IDF routing with audit features; typed Handoff envelopes; "
     "tiered progressive disclosure (L0 name → L1 skill summaries for shortlist → "
     "L2 full contract for chosen agent → L3 tool schemas at invocation), driven "
     "by the router, not the LLM; conformance checks and drift/conflict/"
     "low-confidence detectors"],
    ["4 · Memory", "Beliefs with version control",
     "Mnema: immutable typed assertions, bi-temporality, mandatory provenance, "
     "retraction cascade, sleep/wake consolidation, why()/timeline(); "
     "Memento-style case recall injecting similar past episodes + outcomes at "
     "dispatch"],
    ["5 · Learning & trust", "Learn from failure without touching the flow",
     "Promotion spec for every learnable artifact; per-(domain, kind, model) "
     "scoping so the system is LLM-agnostic; ACE-style delta playbooks; "
     "paired-trajectory auditing as label-free gate evidence; constitution "
     "invariants as hard blocks above all learning"],
], widths=[1.15, 1.7, 3.85])

doc.add_heading("4.4 Architecture: control plane / data plane", level=2)
lead("Data plane (embedded in the customer's runtime)",
     "graxella.mesh / instrument one-line wrap; deterministic router; bounded "
     "multi-hop handoff runtime with loop detection and budgets; healing "
     "dispatch; adapters for LangGraph, CrewAI, AutoGen, MCP, A2A, and plain "
     "callables. No network dependency on the hot path.")
lead("Control plane (multi-tenant service)",
     "Mnema ledger (SQLite at the edge, Postgres at scale); Evidence Gate; "
     "offline miners; audit service (W3C PROV-O export, per-promotion change "
     "reports); value ledger computing ROI and reliability curves as queries. "
     "Horizontally scalable, sharded by namespace and domain; evidence never "
     "leaks across tenants or domains.")

# ===========================================================================
# 5. DIFFERENTIATORS
# ===========================================================================
doc.add_heading("5. Differentiators", level=1)
bullets([
    ("Reliability that compounds with volume", "The only stack where the "
     "reliability curve has a positive slope by mechanism: volume → evidence → "
     "gated promotions → fewer failures. Everyone else's slope is approximately "
     "zero."),
    ("Zero runtime LLM in the substrate", "Routing, healing, governance, and "
     "gate decisions are deterministic and O(1). Costs fall with scale while "
     "LLM-judge approaches grow linearly."),
    ("Audit by construction", "Every decision, outcome, verdict, and promotion "
     "is a cited, immutable assertion. why() is a lookup, not forensics. This "
     "cannot be retrofitted by competitors who never recorded citations."),
    ("Additive adoption, reversible always", "One-line wrap over unchanged "
     "framework code; unwrap anytime and the agents still run. No new syntax, "
     "no user-authored protocol, no rewrite."),
    ("Evidence-gated learning, never silent", "Detection-only governance: "
     "detectors flag loudly; humans and evidence promote; nothing self-corrects "
     "silently. The defined flow is constitutionally protected."),
    ("LLM-agnostic by scoping", "Model ids are recorded on every episode; "
     "promotions are scoped per (domain, kind, model). Swap the LLM and learned "
     "behavior re-validates instead of silently misfiring."),
    ("A benchmarked claim, not a slogan", "The MAST Coverage Scorecard maps "
     "every mechanism to a peer-reviewed failure mode and is re-issued per "
     "release with detection rates from replaying MAST-Data traces."),
])

# ===========================================================================
# 6. COMPETITIVE LANDSCAPE — WHY RELEVANT DESPITE POPULAR SOLUTIONS
# ===========================================================================
doc.add_heading("6. Why Graxella Is Relevant Despite Popular Existing Solutions", level=1)
p("The agent ecosystem is crowded, but every popular category occupies a "
  "different seat. Graxella competes head-on with almost none of them; it "
  "occupies the evidence-substrate seat that all of them leave empty. The "
  "table below is the standing answer to 'doesn't X already do this?'")
table([
    ["Category", "Popular solutions", "What they do", "Why graxella still matters"],
    ["Orchestration frameworks", "LangGraph, LangChain, CrewAI, AutoGen",
     "Author and execute agent workflows",
     "Complements, never competes: graxella wraps them unchanged. Frameworks own "
     "syntax; graxella owns evidence. Framework churn is a tailwind — the ledger "
     "outlives any stack choice."],
    ["Observability / evals", "LangSmith, Langfuse, Arize Phoenix, AgentOps",
     "Trace, visualize, and evaluate runs",
     "They see; they do not learn or govern. Graxella emits standard OTel GenAI "
     "traces INTO them, and adds the layer they lack: decisions with citations, "
     "gated promotions, closed-loop improvement."],
    ["Memory layers", "Mem0, Zep/Graphiti, Letta",
     "Store and retrieve agent memory",
     "They remember facts; mnema governs beliefs — immutable, bi-temporal, "
     "provenance-mandatory, with retraction cascades. ACE's context-collapse "
     "findings argue for exactly this append-only design. Memory here is the "
     "substrate of governance, not a recall cache."],
    ["Guardrails", "NeMo Guardrails, Guardrails AI, LLM firewalls",
     "Static rule checks on inputs/outputs",
     "Rules are written once and rot; graxella's gate is evidence-based and "
     "self-calibrating, with static invariants kept only for the constitution "
     "layer where hard blocks belong."],
    ["MCP gateways / registries", "ContextForge, Kong Agent Gateway, Composio, AWS MCP Gateway",
     "Federate, allowlist, and proxy tool access",
     "Gateways enforce; graxella learns. It rides above them with drift healing, "
     "verified transforms, and live tool trust scores no gateway computes. "
     "Plumbing is delegated, not rebuilt."],
    ["Model routers", "RouteLLM, LiteLLM, Martian",
     "Route queries between cheap and strong models",
     "Graxella learns model-tier assignment from its own outcome ledger per "
     "(domain, kind) and ships it as a gated promotion; LiteLLM is adopted "
     "underneath as provider abstraction."],
    ["Prompt optimizers", "DSPy/GEPA, promptfoo",
     "Compile better prompts offline",
     "Adopted as a component: DSPy is graxella's official compile-time "
     "optimizer. Graxella supplies what DSPy lacks — the trainset (episodes), "
     "the metric (outcomes), and the governance (gate) around deployment."],
    ["Protocols", "MCP, A2A (Linux Foundation), ACP",
     "Standardize tool access and agent messaging",
     "Ridden, not fought: A2A cards and MCP are native surfaces. Peer-reviewed "
     "analysis (arXiv 2606.31498) confirms these protocols cannot express "
     "governance — graxella is the documented missing layer."],
], widths=[1.05, 1.45, 1.35, 2.85])
p("Summary: frameworks build agents, observability watches them, protocols "
  "connect them, guardrails filter them — and none of them makes agents "
  "measurably more reliable next month than this month, with evidence a "
  "regulator can read. That is graxella's seat, and it is empty.", italic=True)

# ===========================================================================
# 7. TARGET AUDIENCE
# ===========================================================================
doc.add_heading("7. Target Audience", level=1)

doc.add_heading("7.1 Ideal customer profile", level=2)
p("Mid-to-large enterprises (1,000+ employees) that have moved past agent "
  "experimentation into (or blocked at the gate of) production, in domains where "
  "auditability is a license to operate. Sweet spot: an AI platform team of "
  "5–50 engineers serving many internal business units, running workflows of "
  "3+ steps across drifting internal APIs, with a compliance function that must "
  "sign off before scale-up.")
bullets([
    "Priority verticals: financial services and insurance (SOX, model risk "
    "management), healthcare (audit-heavy, high blast radius), telecom and "
    "enterprise SaaS operations (high volume, cost pressure), and any EU-exposed "
    "enterprise facing AI Act obligations.",
    "Anti-profile (not now): single-agent chatbot teams, prototype-stage "
    "startups without production volume (no evidence to learn from), and teams "
    "seeking an agent-authoring framework rather than a reliability layer.",
])

doc.add_heading("7.2 Personas", level=2)
table([
    ["Persona", "Role in the deal", "What graxella gives them"],
    ["Head of AI Platform / Platform Engineering", "Economic buyer and champion",
     "The reliability-slope dashboard; governance that scales without headcount; "
     "a control plane that outlives framework churn"],
    ["Staff / Senior AI engineer", "Day-one user, bottom-up adopter",
     "One-line wrap, zero rewrite, loud failures instead of silent ones, "
     "why() instead of log archaeology"],
    ["Compliance / Model Risk officer", "Veto holder turned ally",
     "Cited decision trails, PROV-O export, per-promotion change reports with "
     "approver and rollback point; EU AI Act-grade traceability by construction"],
    ["CFO / Business sponsor", "Renewal decision",
     "Cost per completed task trending down; ROI as a query, not an estimate; "
     "the week-1 vs week-8 pilot report"],
    ["Open-source developer community", "Adoption flywheel and credibility",
     "Apache-2.0 data plane and mnema; standards alignment (OTel GenAI, A2A, "
     "MCP); benchmark publications (MAST scorecard, LongMemEval)"],
], widths=[1.6, 1.5, 3.6])

# ===========================================================================
# 8. SUCCESS METRICS
# ===========================================================================
doc.add_heading("8. Success Metrics (North Stars)", level=1)
table([
    ["Metric", "Definition", "Why it is the headline"],
    ["Reliability slope", "Δ task success rate per 1,000 runs",
     "The compounding claim, falsifiable; everyone else's slope is ~zero"],
    ["Cost per completed task", "Spend ÷ successful completions, trended",
     "The CFO's number; falls as promotions accumulate"],
    ["MTTH (mean time to heal)", "Drift detected → healed in production",
     "Days (human prompt-edit) → minutes (gated auto-promotion)"],
    ["Auto-promotion rate", "% behavior changes shipped by the gate with citations",
     "Governance scaling without headcount"],
    ["Audit answerability", "% decisions with a complete why-chain",
     "100% by construction — the regulator metric"],
    ["MAST detection coverage", "% of taxonomy failure modes detected on replayed traces",
     "The benchmarked reliability claim, re-issued per release"],
], widths=[1.5, 2.3, 2.9])

# ===========================================================================
# 9. ROADMAP
# ===========================================================================
doc.add_heading("9. Roadmap (26 Weeks, Six Phases)", level=1)
p("Each phase ends in buyer-visible value. Detail lives in the Evidence Loop "
  "plan; this is the reference summary.")
table([
    ["Phase", "Weeks", "Delivers", "Exit proof"],
    ["0 · Truth flows", "1–3",
     "Auto-recorded typed outcomes; Memento-style case recall (first visible "
     "value); packaging, tests, loud failures",
     "Every routed task yields a cited decision+outcome pair"],
    ["1 · Evidence Gate", "3–8",
     "Bayesian gate replaces scored policy; one unified proposal pipeline; "
     "paired-trajectory audit evidence",
     "Same proposal blocked cold, auto-approved after N cited successes"],
    ["2 · Handoff runtime", "8–14",
     "Bounded multi-hop dispatch; chain healing; tool trust scores on the "
     "capability graph; MAST's top-2 failure modes addressed",
     "Injected fault at hop 3 self-heals the following week"],
    ["3 · Scale substrate", "14–20",
     "Control/data-plane split; Postgres ledger; OTel GenAI conventions; "
     "A2A card interop; persistent-by-default memory",
     "Route p50 <10ms @ 1k agents; gate p50 <5ms; crash-safe"],
    ["4 · Value ledger", "20–26",
     "ROI as query; reliability-slope dashboard; compliance pack; model-tier "
     "routing savings",
     "The week-1 (71%, $0.48/task) vs week-8 (94%, $0.19/task) pilot report, "
     "all cited"],
    ["5 · Ecosystem", "26+",
     "CrewAI/AutoGen adapters; MCP-native healing; A2A disclosure extension; "
     "evidence federation (transfer_from)",
     "Same workflow crosses frameworks; ledger and curve intact"],
], widths=[1.25, 0.6, 2.75, 2.1])

# ===========================================================================
# 10. GUIDING PRINCIPLES
# ===========================================================================
doc.add_heading("10. Guiding Design Principles", level=1)
bullets([
    ("Compile-time LLM > runtime LLM", "Learning happens offline; dispatch and "
     "governance are deterministic. LLMs propose; they never decide."),
    ("Detection-only governance", "Detectors flag with structured, loud events; "
     "humans and evidence promote. Silent remediation destroys auditability and "
     "is prohibited."),
    ("Memory-grounded, cited, self-calibrating", "Every governance decision "
     "reads from the ledger and writes back to it. No policy engines, no rule "
     "DSLs, no LLM opinions in the loop."),
    ("Silent plumbing, loud failures", "Users author no JSON/YAML/protocol; "
     "peer awareness is automatic; every degradation is a flagged, visible "
     "event; choice is always additive — adding a path never removes one."),
    ("The flow is sacred", "The developer-defined workflow topology is protected "
     "by constitution invariants that no accumulated evidence can override."),
    ("Ride standards, own the evidence", "OTel GenAI, A2A, MCP, PROV-O are "
     "adopted, not reinvented. The evidence loop is the only proprietary thing, "
     "and it is the moat."),
])

# ===========================================================================
# 11. SCOPE
# ===========================================================================
doc.add_heading("11. Scope", level=1)
doc.add_heading("11.1 In scope", level=2)
bullets([
    "The reliability/governance control plane and embedded data plane described "
    "in Section 4.4, across the five capability stages.",
    "Adapters for LangGraph/LangChain (first), CrewAI and AutoGen (Phase 5), "
    "MCP tools, A2A remote agents, and plain callables.",
    "Mnema as both the graxella substrate and a standalone open-source memory "
    "engine (dual life is intentional).",
    "The value ledger, compliance pack, and MAST regression suite.",
])
doc.add_heading("11.2 Out of scope", level=2)
bullets([
    "Authoring frameworks: graxella will not invent an agent-definition syntax "
    "or compete with LangGraph/CrewAI for workflow authorship.",
    "Model serving/hosting, GPU infrastructure, fine-tuning pipelines.",
    "Gateway plumbing (federation, authN/Z proxying) — delegated to the MCP "
    "gateway ecosystem; graxella integrates above it.",
    "Autonomous flow rewriting: the system never restructures the user's "
    "workflow topology, by principle, not just by phase.",
])

# ===========================================================================
# 12. COMPONENT INVENTORY
# ===========================================================================
doc.add_heading("12. Component Inventory", level=1)
table([
    ["Component", "Origin", "Role in graxella", "Status (Aug 2026)"],
    ["graxella core", "Umbrella package",
     "mesh/instrument wrap, Society routing facade, gate, constitution, tracer, "
     "healing dispatch, agenda miners, PROV-O audit, CLI",
     "Phases 1–7 code-complete; Phase 0 hardening owed (tests, packaging, "
     "closed loop)"],
    ["mnema (B12)", "Memory intelligence block",
     "Immutable belief ledger: typed assertions, bi-temporality, provenance, "
     "retraction cascade, consolidation, why()/timeline()",
     "Strongest component; property-tested, CI, LongMemEval harness"],
    ["agent2society (B11)", "A2A layer block",
     "Deterministic routing with audit features, Handoff envelope, conformance, "
     "governance detectors, explanation store, backtest-gated optimizer",
     "Shipped with 15-file test suite; vendored under graxella"],
    ["axon tool fabric (B8 + B10)", "MCP orchestration + self-healing blocks",
     "Tool knowledge graph and semantic discovery (B8); drift interception, "
     "healer→compiler→verifier transforms, federated registry, sentinel "
     "(canary/shadow/differ/forecaster), EU AI Act policy pack (B10)",
     "B10 production-shaped, to be integrated behind the gate; B8 archived as "
     "R&D, concepts ported"],
    ["Evidence Gate", "New (Phase 1)",
     "Bayesian promotion decisions from ledger priors per (domain, kind, model); "
     "replaces the interim scored GatePolicy",
     "Designed (see gate design note); the current scored implementation is "
     "explicitly interim"],
], widths=[1.35, 1.35, 2.9, 1.7])

# ===========================================================================
# 13. RISKS
# ===========================================================================
doc.add_heading("13. Risks and Mitigations", level=1)
table([
    ["Risk", "Impact", "Mitigation"],
    ["Cold start: no evidence, no value", "New deployments see governance "
     "friction before compounding kicks in",
     "Phase 0 case recall gives value from week one; observe-only adoption step "
     "carries zero risk; transfer_from seeds warm priors from sibling domains"],
    ["Evidence poisoning / gaming", "Attacker or flaky agent floods ledger with "
     "fake positives to loosen gates",
     "Provenance-diversity floor (≥K independent sessions/operators); "
     "constitution hard blocks are never evidence-overridable; immutable ledger "
     "makes tampering visible"],
    ["Framework API churn breaks adapters", "Duck-typed integration points "
     "(e.g., LangGraph internals) drift",
     "Adapter contract tests pinned per framework version in CI (Phase 0); "
     "adapters are the only churn surface by design"],
    ["Category confusion ('another framework?')", "Positioning collapses into "
     "the crowded orchestration shelf",
     "Discipline in Section 6's seat map; never ship an authoring syntax; lead "
     "with control-plane language and the reliability slope"],
    ["Scored-gate interim ossifies", "The rejected GatePolicy design gets "
     "mistaken for the product",
     "Charter and code mark it interim; Phase 1 replaces it; the gate design "
     "note is the binding spec"],
    ["Scale substrate arrives late", "A successful pilot dies at the scale-up "
     "meeting on SQLite/JSONL",
     "Phase 3 is scheduled before the value-ledger sales motion (Phase 4); "
     "perf targets are exit criteria, not aspirations"],
    ["Solo-founder bandwidth", "Six phases across five components is a "
     "platform-company workload",
     "Phase discipline (each stage reaches demo-with-evidence before the next); "
     "adopt-not-build shortlist (DSPy, LiteLLM, OTel, A2A SDK, Graphiti-watch) "
     "keeps surface area honest"],
], widths=[1.7, 2.0, 3.0])

# ===========================================================================
# 14. REFERENCES
# ===========================================================================
doc.add_heading("14. Reference Documents", level=1)
bullets([
    "The Evidence Loop — program strategy and 26-week plan (artifact: "
    "claude.ai/code/artifact/bd7ac6a1-4588-4c6d-8d77-b47b66e5881f)",
    "Graxella Research Radar — papers/OSS adopt-borrow-watch sweep (artifact: "
    "claude.ai/code/artifact/603d25fa-95e4-4485-b7c4-3b50f50bc97f)",
    "MAST Coverage Scorecard — graxella vs the 14 failure modes (artifact: "
    "claude.ai/code/artifact/2aafe16f-f501-48ec-8a31-e7cd3ee0d4f2)",
    "Graxella_Flagship_Deck.pptx — 18-slide program deck (this folder)",
    "Key external references: MAST (arXiv 2503.13657); Governance Gaps in Agent "
    "Interoperability Protocols (arXiv 2606.31498); ACE — Agentic Context "
    "Engineering (arXiv 2510.04618); Memento (arXiv 2508.16153); GEPA "
    "(github.com/gepa-ai/gepa); OTel GenAI semantic conventions "
    "(opentelemetry.io); A2A v1.0.1 (Linux Foundation)",
])

# ===========================================================================
# 15. GLOSSARY
# ===========================================================================
doc.add_heading("15. Glossary", level=1)
table([
    ["Term", "Meaning"],
    ["Evidence Loop", "The six-stage cycle (dispatch → outcome → ledger → "
     "proposal → gate → promotion) by which volume becomes reliability"],
    ["Promotion", "A versioned, gated, rollbackable behavior change carrying "
     "citations — the unit of improvement"],
    ["Ledger", "Mnema's immutable store of typed assertions (decisions, "
     "outcomes, verdicts) with mandatory provenance"],
    ["Evidence Gate", "Bayesian promotion decision from cited priors per "
     "(domain, kind, model); zero LLM; self-calibrating thresholds"],
    ["Constitution", "Hard invariants (flow topology, budgets, forbidden "
     "actions) that no evidence can override; detection-only enforcement"],
    ["Optimistic orchestration", "Doctrine: the defined flow always runs as "
     "written; learning is background, additive, and reversible"],
    ["Waiting tokens", "Tokens burned by supervisor LLMs deliberating over "
     "agent selection — eliminated by deterministic routing"],
    ["Progressive disclosure (L0–L3)", "Tiered capability revelation driven by "
     "the router: name → skill summaries → full contract → tool schemas"],
    ["Case recall", "Injecting top-k similar past episodes with outcomes into "
     "the dispatched agent's context (Memento pattern)"],
    ["Blast radius", "Scope of a proposed change; wide blast requires "
     "overwhelming same-tuple evidence or human sign-off"],
    ["MTTH", "Mean time to heal: drift detected → healed in production"],
    ["Sentinel", "Proactive drift detection: canary calls, shadow traffic, "
     "schema diffing, drift forecasting (from B10)"],
], widths=[1.9, 4.8], header=True)

# ---------------------------------------------------------------------------
out = Path(__file__).parent / "Graxella_Project_Charter.docx"
doc.save(out)
print(f"Wrote {out}")
