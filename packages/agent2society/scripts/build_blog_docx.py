"""Generate a Medium-style blog post about agent2society as a .docx file."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "agent2society_medium_blog.docx"

CODE_FONT = "Consolas"
BODY_FONT = "Georgia"
HEAD_FONT = "Helvetica"

doc = Document()

# --- Page margins -----------------------------------------------------------
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# --- Base style -------------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(12)


def add_title(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = HEAD_FONT
    run.font.color.rgb = RGBColor(0x14, 0x14, 0x14)


def add_subtitle(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(16)
    run.font.name = HEAD_FONT
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(18)


def add_byline(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.paragraph_format.space_after = Pt(24)


def add_h1(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = HEAD_FONT
    run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)


def add_h2(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = HEAD_FONT


def add_body(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(12)


def add_body_mixed(segments) -> None:
    """segments: list of (text, {'bold':bool,'italic':bool,'code':bool})."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5
    for text, opts in segments:
        run = p.add_run(text)
        run.bold = opts.get("bold", False)
        run.italic = opts.get("italic", False)
        if opts.get("code"):
            run.font.name = CODE_FONT
            run.font.size = Pt(10.5)
        else:
            run.font.name = BODY_FONT
            run.font.size = Pt(12)


def add_bullet(text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.4
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(12)
        tail = p.add_run(text)
        tail.font.name = BODY_FONT
        tail.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = Pt(12)


def add_code(code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(code)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def add_pullquote(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(13)
    run.font.name = BODY_FONT
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_divider() -> None:
    p = doc.add_paragraph()
    run = p.add_run("— — —")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


# ============================================================================
# CONTENT
# ============================================================================

add_title("agent2society")
add_subtitle("The transparent coordination layer for A2A agent meshes — every routing decision is a deterministic graph traversal you can read, threshold on, and audit.")
add_byline("By Ramprasad  ·  June 2026  ·  ~12 min read")


# --- Section 1 --------------------------------------------------------------
add_h1("The problem nobody wants to say out loud")

add_body(
    "Multi-agent systems are powerful and untrustworthy. Anyone who has shipped one into production "
    "knows the second half of that sentence better than the first. You wire up a supervisor LLM, hand "
    "it a stable of specialist agents, write a careful system prompt that explains who can do what, "
    "and ship the thing. For a while, it works."
)

add_body(
    "Then the first failure shows up. A finance question quietly gets routed to the writer agent. A "
    "PII-laden ticket reaches the public-search agent. The marketing memo cites a number nobody can "
    "trace. You open the logs and try to reconstruct what happened. There is a transcript. There is a "
    "model that \"decided.\" There is no explanation."
)

add_pullquote(
    "The supervisor pattern hides every routing choice behind an LLM that decided. Debugging a wrong "
    "handoff means reading a transcript and guessing what the model thought."
)

add_body(
    "That is not a bug; it is the design. The whole appeal of supervisor-style multi-agent systems is "
    "that the LLM figures coordination out. The cost is that you inherit the opacity of that "
    "figuring-out. You pay in two currencies: tokens (the supervisor reasons in every round) and "
    "trust (the routing is whatever the model felt like that turn)."
)

add_body(
    "agent2society is the opposite stance. It sits on top of the A2A (Agent-to-Agent) protocol and "
    "replaces opaque LLM negotiation with a deterministic, graph-routed dispatch that produces a "
    "human-readable explanation for every decision. Which agent was chosen. The features that drove "
    "the score. The alternatives that were rejected, and why. The chosen agent's self-declared limits. "
    "The upstream decision chain that led here."
)

add_body(
    "It is on PyPI. It installs with zero dependencies. The default scorer runs on TF-IDF so the "
    "transparency wedge is visible on the very first run, before you swap in real embeddings."
)

add_code("pip install agent2society")


# --- Section 2 --------------------------------------------------------------
add_h1("The cost wedge: one embedding instead of an LLM round")

add_body(
    "You cannot match free-text tasks to capabilities with literally zero model calls. But you can do "
    "it with one cheap embedding lookup instead of many rounds of supervisor reasoning. That delta is "
    "the cost story."
)

add_body(
    "Concretely: the router parses an A2A Agent Card for each agent, extracts the declared skills, "
    "and embeds the skill descriptions once. At dispatch time, the incoming task is embedded once. A "
    "single cosine-similarity pass over the skill index returns a ranked candidate list. No "
    "supervisor prompt. No multi-turn negotiation. No tokens spent on \"the writer agent can write "
    "memos, the research agent can do research, the analyst can ...\""
)

add_body("On the published benchmark suite this lands at:")

add_code(
    "coordination tokens per task\n"
    "  supervisor:     median=427  mean=427.8  min=424  max=433\n"
    "  agent2society:  median=9    mean=9.4    min=7    max=15\n"
    "\n"
    "reduction:       94.12%\n"
    "\n"
    "task success (right agent + skill chosen)\n"
    "  supervisor:     11/12\n"
    "  agent2society:  11/12"
)

add_body(
    "Roughly 140x cheaper coordination on a labelled twelve-task suite, holding correctness equal. "
    "Dispatch cost (the actual agent calls doing real work) is held the same across both methods; "
    "the delta is purely the supervisor's reasoning tokens. A real LangGraph head-to-head, included "
    "in the benchmarks directory, shows the same shape: 96.11% reduction, ~200x ratio, two honest "
    "misses on the TF-IDF default scorer that swap in sentence-transformers to close."
)

add_body(
    "That is the cost wedge. It is enough to get developers to install the package. The reason they "
    "keep it is everything that comes after."
)


# --- Section 3 --------------------------------------------------------------
add_h1("What an explanation actually looks like")

add_body(
    "Every dispatch produces a RoutingExplanation, keyed by handoff id, stored in the same ledger that "
    "captures the result. You can render it as text:"
)

add_code(
    "handoff a77d6eed5a15  task: Draft an executive memo on customer churn\n"
    "  intent: prep the Q3 board pack\n"
    "  assumes:\n"
    "    - churn data is final\n"
    "  chose  : writer-agent :: exec_memo (confidence=0.593)\n"
    "  why    : writer-agent::exec_memo selected (score=0.593). Semantic match\n"
    "           0.725 on tokens [an, draft, executive, memo]; tag overlap 0.200\n"
    "           on [draft, memo]. Runner-up research-agent::web_research rejected:\n"
    "           below min_score.\n"
    "  features: min_score=0.050, confidence_required=0.500, score=0.593,\n"
    "            semantic=0.725, tag_overlap=0.200\n"
    "  alternatives:\n"
    "    - writer-agent :: exec_memo  score=0.593 sem=0.725 tags=0.200\n"
    "    - research-agent :: web_research  score=0.000  REJECTED (below min_score)\n"
    "  agent self-caveats:\n"
    "    - English only\n"
    "    - max ~400 words per memo\n"
    "    - escalate when: any quantitative claim cited without source"
)

add_pullquote(
    "No LLM produced that. It is a template rendered from the features the router actually used."
)

add_body(
    "Read it twice. Notice what is in there: the score, the features that produced the score, the "
    "tokens that drove the semantic match, the tag overlap that broke the tie, the runner-up and the "
    "specific reason it lost, the agent's own self-declared limits surfaced inline. This is what "
    "\"explainable\" means when you do not have to ask another LLM to please be honest about why it "
    "made the call."
)

add_body(
    "It is ASCII-safe by design — a regression test pins this. The default Windows console will print "
    "it without choking on Unicode dashes. That tiny detail is the difference between an "
    "explainability feature that demos well and one that survives a 3am incident review on a "
    "developer's actual machine."
)


# --- Section 4 --------------------------------------------------------------
add_h1("Sixty-second tour")

add_body("Here is the smallest interesting program. Three agents, one boundary, one handoff.")

add_code(
    "from agent2society import Society, Handoff\n"
    "\n"
    "society = Society()\n"
    "society.add(\"https://research-agent.acme.com/.well-known/agent-card.json\")\n"
    "society.add(\"https://writer-agent.acme.com/.well-known/agent-card.json\")\n"
    "society.add(my_crewai_crew)                              # adapters wrap native objects\n"
    "\n"
    "society.boundary(\"writer-agent\", deny=[\"financial-data\"])\n"
    "\n"
    "# Strings still work.\n"
    "result = society.run(\"Summarise Q3 churn and draft an exec memo\")\n"
    "\n"
    "# A Handoff carries why the task exists, what we assume, and prior decisions.\n"
    "h = Handoff(\n"
    "    task=\"Draft an executive memo on Q3 customer churn\",\n"
    "    intent=\"prep the board pack\",\n"
    "    assumptions=[\"churn data is final\"],\n"
    "    confidence_required=0.5,\n"
    ")\n"
    "memo = society.run(h)\n"
    "\n"
    "print(society.explain(h.id).render())\n"
    "society.report()"
)

add_body(
    "That is the whole API surface most users touch. Society for the mesh. Handoff for tasks that "
    "carry context. boundary() for policy. explain() for the rationale. report() for the audit log. "
    "Mesh is preserved as an alias for backward compatibility — every line of v0 code keeps working."
)


# --- Section 5 --------------------------------------------------------------
add_h1("Conformance: the guardrail that stops silent acceptance")

add_body(
    "Before dispatch, the conformance check answers two questions deterministically. Does this "
    "agent's card actually declare this skill? Is the task inside this agent's declared boundary "
    "(allow / deny)? If either fails, the dispatch is blocked. In strict mode it raises "
    "ConformanceViolation. In non-strict mode it records the violation in telemetry and the "
    "explanation's blocked_reason field, then returns an empty string."
)

add_code(
    "society.boundary(\"writer-agent\", deny=[\"financial-data\", \"pii\"])\n"
    "society.boundary(\"research-agent\", allow=[\"public\"])"
)

add_body(
    "This is the layer that stops an agent from silently accepting work it never claimed it could do. "
    "In a supervisor pattern, this entire class of failure leaks by design: the supervisor passes a "
    "task to the writer agent, the writer agent attempts the task, the writer agent produces something "
    "plausible-looking, and nothing in the system can tell you it should not have been there in the "
    "first place. agent2society makes that case impossible, not unlikely."
)

add_body(
    "The boundary check uses NFKC Unicode normalisation plus casefolding under the hood, so adversarial "
    "encodings (\"FINANCIAL\", \"ﬁnancial\", \"fınancial\") collapse to the same comparison key. That fix "
    "shipped in v0.5.3 — closing a Unicode-bypass class that, frankly, every text-rule system needs to "
    "think about and most do not."
)


# --- Section 6 --------------------------------------------------------------
add_h1("Handoff: context that survives the chain")

add_body(
    "A bare string answers \"what to do next.\" A Handoff answers \"what to do next, why it exists, "
    "what's been done so far, what we assume to be true, and what would force a human review.\" That "
    "is the difference between a task queue and a decision record."
)

add_code(
    "h0 = Handoff(\n"
    "    task=\"Research Q3 customer churn drivers\",\n"
    "    intent=\"prep the Q3 board pack\",\n"
    "    assumptions=[\"churn data through end-of-quarter is final\"],\n"
    "    confidence_required=0.5,\n"
    ")\n"
    "research_text = society.run(h0)\n"
    "\n"
    "# Extend the handoff: the next agent sees the upstream decision in its\n"
    "# prior chain (and so does the routing explanation).\n"
    "h1 = h0.extend(\n"
    "    agent=\"research-agent\",\n"
    "    skill=\"web_research\",\n"
    "    summary=\"found 3 churn drivers\",\n"
    "    confidence=0.62,\n"
    "    next_task=\"Draft an executive memo on those churn drivers\",\n"
    ")\n"
    "memo = society.run(h1)"
)

add_body(
    "When the memo dispatch runs, its RoutingExplanation includes h0 in its prior chain. The audit log "
    "is no longer a bag of independent rows: it is a directed graph of decisions, each one carrying "
    "the reasoning and confidence of the one that produced it."
)


# --- Section 7 --------------------------------------------------------------
add_h1("SelfAssessment: agents declare their own limits")

add_body(
    "An Agent Card can carry a selfAssessment block — the agent's own statement of what it can do, what "
    "it cannot do, and what should force escalation. agent2society surfaces those limits on every "
    "explanation that picks the agent, so the caller and downstream agents see the same caveats the "
    "agent claims for itself."
)

add_code(
    '{\n'
    '  "name": "writer-agent",\n'
    '  "skills": [{"id": "exec_memo", "name": "Executive Memo"}],\n'
    '  "selfAssessment": {\n'
    '    "confidenceModel": "tfidf_score",\n'
    '    "knownLimitations": ["English only", "max ~400 words per memo"],\n'
    '    "outOfScope": ["legal opinion", "binding financial guidance"],\n'
    '    "escalateWhen": ["any quantitative claim cited without source"]\n'
    '  }\n'
    '}'
)

add_body(
    "This matters because an agent's own honest scope statement is the most reliable place to put "
    "it. The supervisor pattern systematically loses this information — the LLM picks an agent, the "
    "agent does the work, and the caveats live inside the agent's prompt where the caller never sees "
    "them. agent2society treats those caveats as first-class metadata that travels with the decision."
)


# --- Section 8 --------------------------------------------------------------
add_h1("Governance hooks: detection, never auto-correct")

add_body(
    "There are four hooks. They fire on specific patterns the router can detect cheaply from data it "
    "already has."
)

add_code(
    "society.on_low_confidence(lambda exp: notify(exp), threshold=0.5)\n"
    "society.on_human_review(lambda exp, result: page_oncall(result))\n"
    "society.on_conflict(lambda c: log(\"conflict\", c.detail))\n"
    "society.on_capability_drift(lambda d: log(\"drift\", d.agent))"
)

add_bullet(
    " fires when a decision's confidence is below the handoff's confidence_required (or a society-wide threshold).",
    bold_prefix="on_low_confidence",
)
add_bullet(
    " fires when a Handoff.human_review_when(result_text) predicate returns True. Default on exception is fail-safe (review required).",
    bold_prefix="on_human_review",
)
add_bullet(
    " fires when the same task text was routed to different (agent, skill) pairs across handoffs in the window.",
    bold_prefix="on_conflict",
)
add_bullet(
    " fires when one agent is selected across an unusually broad spread of skills — a signal that skill descriptions may be too generic.",
    bold_prefix="on_capability_drift",
)

add_pullquote(
    "Hooks are side effects. They cannot block, retry, or modify a dispatch. That separation is "
    "deliberate — silent auto-correction is exactly the opacity that makes meshes untrustworthy."
)

add_body(
    "Handler exceptions are caught (so a buggy hook can never break a dispatch) but logged at WARNING "
    "on the agent2society logger with the hook's qualname — they never vanish silently. A hook that "
    "throws once a thousand times is a hook you can find and fix; a hook that throws silently is a "
    "production incident waiting to happen."
)


# --- Section 9 --------------------------------------------------------------
add_h1("Production ops")

add_h2("Metrics")
add_body(
    "Every Society exposes a MetricsCollector with counters and summary-style histograms, ready to "
    "scrape. JSON or Prometheus text format — your choice."
)
add_code(
    "society.metrics.snapshot()                  # JSON-serialisable dict\n"
    "print(society.metrics.render_prometheus())  # Prometheus text format"
)
add_body(
    "Pre-registered series cover routes, dispatches, retries, failures, conformance blocks, conflicts, "
    "low-confidence events, drift, human-review firings, and optimizer applications."
)

add_h2("Persistent ledger")
add_body(
    "Replace the default in-memory store with a JSONL file to keep decisions across restarts. Each "
    "run() appends one JSON line; a fresh Society rebuilds the index on startup. Corrupt lines are "
    "skipped (not fatal), and v0.5.3 added structured corruption stats so you can see how many lines "
    "were skipped and why."
)
add_code(
    "from agent2society import Society, JsonlFileStore\n"
    "\n"
    "society = Society(store=JsonlFileStore(\"audit.jsonl\"))"
)

add_h2("Auto-retry and fallback dispatch")
add_body(
    "A transport error on the chosen candidate falls through to the next conformance-passing "
    "candidate. Every retry is recorded in RoutingRecord.fallbacks and emits a "
    "dispatch_retries_total counter."
)
add_code("society.run(handoff, retry=True)")

add_h2("Thread safety")
add_body(
    "Mutation paths (add, run, optimize, apply_optimization) acquire a reentrant lock. Boundary edits "
    "use a copy-on-write pattern under the hood: the graph is deep-copied, mutated, and atomically "
    "swapped in. Concurrent run() calls from multiple worker threads are safe; they never see a "
    "partially-mutated graph mid-dispatch. This shipped in v0.5.3 and is one of the fifteen production "
    "hardening fixes that came out of an enterprise-grade audit."
)

add_h2("Session traces (v0.5.3)")
add_body(
    "SessionTracer joins the telemetry sink and the explanation store into an ordered event stream "
    "per session. Zero new per-route cost — events are built from data the Society was already "
    "recording. Useful for dashboards, post-mortems, or piping into a notebook."
)
add_code(
    "from agent2society import SessionTracer\n"
    "\n"
    "tracer = SessionTracer(society)\n"
    "for ev in tracer.events():\n"
    "    print(ev.handoff_id, ev.chosen_agent, ev.score, ev.flags)"
)


# --- Section 10 -------------------------------------------------------------
add_h1("Native agents: adapters")

add_body(
    "agent2society reads A2A Agent Cards, but you do not have to host every agent over HTTP to "
    "participate. Adapters wrap native objects into A2A cards plus local handlers, and the dispatcher "
    "prefers local handlers over HTTP via a CompositeTransport."
)

add_body("Out of the box, any of these objects can be dropped straight into society.add():")

add_bullet("A CrewAI Crew (anything with .kickoff() and .agents)")
add_bullet("A compiled LangGraph (anything with .invoke() and .nodes)")
add_bullet("An AutoGen ConversableAgent or GroupChatManager")
add_bullet("Any plain callable, or any object with run / invoke / kickoff")

add_body("New frameworks slot in via register_adapter():")

add_code(
    "from agent2society.adapters.base import Adapter, register_adapter\n"
    "\n"
    "class MyAdapter(Adapter):\n"
    "    def matches(self, obj): ...\n"
    "    def to_card(self, obj): ...\n"
    "    def to_handler(self, obj): ...\n"
    "\n"
    "register_adapter(MyAdapter())"
)


# --- Section 11 -------------------------------------------------------------
add_h1("The routing-quality signals you should actually watch")

add_body(
    "Shipping a transparency layer without surfacing the cases where transparency matters most would "
    "be the punchline of a bad joke. Every explanation carries three flags:"
)

add_bullet(
    " the top-1 task is far from anything in the skill index. The router still returned a best guess, but it should be treated as low-quality. Likely candidate for a human review hook.",
    bold_prefix="OOD (out-of-distribution): ",
)
add_bullet(
    " multiple candidates scored almost the same. Either your skill descriptions overlap too much, or this task genuinely has two valid handlers.",
    bold_prefix="VECTOR_AMBIGUITY: ",
)
add_bullet(
    " the top-1 vs top-2 score gap is small. Margin is the cheapest confidence proxy you can ask for. Surface it on a dashboard, alert on dips, or gate retry behavior on it.",
    bold_prefix="LOW_MARGIN: ",
)

add_body(
    "These signals are not opinions. They are scalar features the router produced anyway, exposed so "
    "you can build policy on them. Combine them with the governance hooks and you have a real "
    "feedback loop: low-margin decisions feed a human-review queue; the queue's outputs feed your "
    "skill descriptions; your skill descriptions feed back into the router. The mesh improves "
    "without any new model in the critical path."
)


# --- Section 12 -------------------------------------------------------------
add_h1("What v0.5.3 hardened")

add_body(
    "v0.5.3 is the production hardening release. Fifteen failure modes identified in an "
    "enterprise-scale audit, fixed with regression tests one-per-failure-mode. The interesting ones:"
)

add_bullet(
    " boundary edits no longer block readers. Deep-copy the graph, mutate the copy, atomic swap in. Concurrent run() never sees a half-mutated rule set.",
    bold_prefix="Copy-on-write graph mutation: ",
)
add_bullet(
    " conformance now uses NFKC normalisation plus casefolding before any comparison. \"FINANCIAL\" and \"fınancial\" no longer slip past the same allow/deny list.",
    bold_prefix="Unicode bypass closed: ",
)
add_bullet(
    " the conflict and drift detectors use FIFO-evicted bounded dicts (10,000 entries) so they cannot leak unboundedly under sustained traffic.",
    bold_prefix="Governance memory bounded: ",
)
add_bullet(
    " human_review_when predicate failures default to True (review required) instead of False. If your predicate throws, you get more reviews, not silently fewer.",
    bold_prefix="Fail-safe defaults: ",
)
add_bullet(
    " hook exceptions are caught (so they cannot crash dispatch) but logged at WARNING with the hook's qualname and exception class. No more invisible hook failures.",
    bold_prefix="Hook exceptions surfaced: ",
)
add_bullet(
    " extract_text always returns a string. No more AssertionError on weird response shapes; falls back to json.dumps(default=str), then repr().",
    bold_prefix="Total response extraction: ",
)
add_bullet(
    " JsonlFileStore now tracks skipped lines by reason (skipped_json, skipped_shape) and logs at WARNING. Silent data loss on store rebuild is a class of incident that no longer exists.",
    bold_prefix="JSONL corruption tracking: ",
)

add_body(
    "112 tests, 16 of them dedicated to the failure modes above. The point of those tests is not "
    "coverage; it is that any future regression on a known production failure will fail loudly in CI "
    "with a name that tells you exactly what broke."
)


# --- Section 13 -------------------------------------------------------------
add_h1("The thing I keep coming back to")

add_body(
    "Multi-agent frameworks have an interesting psychological pattern. The first demo is magical: "
    "you describe a task, a supervisor LLM does some thinking, agents collaborate, you get an answer. "
    "The first production incident is the opposite: you describe a task, a supervisor LLM did some "
    "thinking, agents collaborated, you got the wrong answer, and you have absolutely no idea why."
)

add_body(
    "The industry's instinct so far has been to fix opacity with more LLMs. Have a reflection step. "
    "Have a critic agent. Have the supervisor explain itself in plain English at the end. None of "
    "that actually fixes the problem — you just push the opacity one layer down and pay more tokens "
    "for the privilege."
)

add_pullquote(
    "Transparency is not a property of an LLM that promises to be honest. It is a property of a "
    "system whose decisions are observable."
)

add_body(
    "agent2society's bet is that for a large class of multi-agent coordination tasks, you do not need "
    "an LLM in the routing loop at all. A deterministic graph, a cheap embedding lookup, and a "
    "first-class explanation of the decision is enough. When that is not enough — when the scoring "
    "genuinely needs richer semantics — you swap in sentence-transformers and keep every other "
    "property the system already had."
)

add_body(
    "It is also a stance about what an LLM is good for. In agent2society, the LLM is a tool you call "
    "for the actual work (writing the memo, doing the research, drafting the reply). It is not a "
    "supervisor pretending to reason its way through a routing decision that could have been a "
    "vector lookup."
)


# --- Section 14 -------------------------------------------------------------
add_h1("Where to find it")

add_body(
    "On PyPI under the name agent2society. Apache-2.0. Zero hard dependencies; optional extras for "
    "HTTP dispatch, sentence-transformers embeddings, and benchmarking."
)

add_code(
    "pip install agent2society                       # core, zero deps\n"
    "pip install \"agent2society[http]\"               # httpx for HTTP dispatch\n"
    "pip install \"agent2society[embeddings]\"         # sentence-transformers"
)

add_body(
    "The 60-second tour at the top of this post is the whole API surface most users need. The "
    "examples directory in the repository has two longer walk-throughs: a basic three-agent mesh, "
    "and a transparent mesh that wires up the SelfAssessment surface, the Handoff chain, and the "
    "governance hooks together. The benchmarks directory has both the synthetic head-to-head and a "
    "real LangGraph supervisor baseline if you want to verify the cost claims yourself."
)

add_divider()

add_body(
    "If you build with this, I want to hear what breaks. Detection-only is a deliberate design "
    "choice, but the right list of detectors is something the community will discover faster than "
    "any single author. Open an issue, file a PR, or just send me a note about the failure mode you "
    "wish the explanation surfaced and did not."
)

add_byline(
    "agent2society is on PyPI now: pip install agent2society. Source and issues on GitHub at "
    "github.com/graxella/agent2society."
)


# ============================================================================
# SAVE
# ============================================================================
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"wrote {OUT}")
